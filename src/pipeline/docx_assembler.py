"""SOP Factory 2 docx assembler (CLAUDE.md: "extend it, do not rewrite it").
Drives the existing SOPBuilder engine (python-docx-based, VRSI-formatted,
at C:\\Users\\Brian\\Documents\\SOP_Factory_2\\template\\sop_lib.py) from a
manifest plus already-rendered step text (task-06/12) and already-annotated
screenshots (task-10). This module calls the engine's public methods in a
new order specific to a captured SOP session, plus a thin repo-side layer
(docx_fields.py) that turns the engine's own plain-text TOC into a real
Word TOC field without touching the external, unversioned engine file."""

import os
import sys
from pathlib import Path

from docx.shared import RGBColor

from pipeline.assembler import step_heading, toc_lines
from pipeline.claim_coverage import parse_verify_line
from pipeline.docx_fields import add_toc_field, enable_update_fields_on_open, set_outline_level
from pipeline.resource_path import resource_path

DEFAULT_SOP_FACTORY_2_DIR = Path(r"C:\Users\Brian\Documents\SOP_Factory_2\template")


def sop_factory_2_dir():
    """Dev/test mode: the external SOP_Factory_2 clone (never vendored into
    this repo — see the module docstring). Frozen mode: task-10's
    PyInstaller spec bundles just the engine module + template assets
    (not the whole external working project) under "sop_factory_2" inside
    the frozen bundle, so resolution there goes through resource_path()
    like every other in-repo resource. An env var override always wins,
    in either mode, for pointing at a non-default clone location."""
    override = os.environ.get("SOPFORGE_SOP_FACTORY_2_DIR")
    if override:
        return Path(override)
    if getattr(sys, "frozen", False):
        return resource_path("sop_factory_2")
    return DEFAULT_SOP_FACTORY_2_DIR


def _import_sop_builder():
    template_dir = str(sop_factory_2_dir())
    if template_dir not in sys.path:
        sys.path.insert(0, template_dir)
    from sop_lib import SOPBuilder

    return SOPBuilder


def _narrative_body(sop, narrative_text):
    """Writes narrative_text as one or more paragraphs, rendering any
    [verify]-flagged line (claim_coverage.parse_verify_line, which reverses
    render_verify_blockquote's own format -- the shared single point of
    truth export_pdf.py's counterpart also uses) as a distinct styled
    callout instead of raw "> [verify] (claim-...)" text that reads as
    debug scaffolding in a shipped document. The claim id itself is dropped
    from what's *shown* — it stays meaningful in the sidecar report, not
    the reader-facing doc."""
    for line in narrative_text.splitlines():
        claim_text = parse_verify_line(line)
        if claim_text is not None:
            p = sop.bullet_rich([("Needs verification: ", True), (claim_text, False)])
            label_run = p.runs[1]  # runs[0] is the bullet_rich "•  " marker
            label_run.italic = True
            label_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif line.strip():
            sop.paragraph(line)


def _step_bullet(sop, step, result):
    """A step's own generated text as a bullet, with its target element's
    name bolded inline (bullet_rich) when that name appears verbatim in the
    text -- e.g. "Click **Save** in the SmartDeploy Console window." Falls
    back to a plain bullet whenever there's no element name, or the
    generated text doesn't literally contain it (an LLM reply that
    paraphrased instead of quoting), since bolding a substring that isn't
    actually there would either crash or silently do nothing useful."""
    text = result["text"]
    name = step.element.name
    if name and name in text:
        before, _, after = text.partition(name)
        sop.bullet_rich([(before, False), (name, True), (after, False)])
    else:
        sop.bullet(text)


def assemble_docx(
    manifest,
    step_results,
    annotated_dir,
    output_path,
    revision="1.0",
    date="01/01/2026",
    author="SOPForge",
    doc_no=None,
    narrative_text=None,
):
    """Builds a complete docx from a manifest's steps (already rendered via
    task-06/task-12's step_results, one dict per step with a "text" key)
    and their annotated screenshots (already written by task-10's
    annotate_click into annotated_dir, one file per step named
    step.screenshot). Returns (output_path, warnings) — warnings is
    SOPBuilder's own missing-image/unpatched-header list.

    `date`/`author`/`doc_no` are meant to be the caller's real, computed
    values (assembler.py's format_doc_date/doc_number, config.py's
    DocumentConfig) -- the defaults here exist only for direct/test callers
    that don't care, never as production fallbacks."""
    sop_builder_cls = _import_sop_builder()
    factory_dir = sop_factory_2_dir()

    sop = sop_builder_cls(
        template_docx=factory_dir / "SOP_TEMPLATE_WITH_PHOTOS.docx",
        output_docx=output_path,
        active_dir=annotated_dir,
        revision=revision,
        date=date,
    )
    title = manifest.session.title or manifest.session.id
    sop.title_page(title.upper(), author=author, doc_no=doc_no)

    # A real Word TOC field (docx_fields.py), not SOPBuilder's own
    # plain-text toc() -- the cached text is the same toc_lines() would
    # show either way, but Word can now actually build/update it. Every
    # heading below is marked with a matching outline level (0 for section
    # headings, 1 for step headings) since SOPBuilder's heading1/heading2
    # carry no named Word style for the field's \u switch to key off.
    add_toc_field(sop, toc_lines(manifest, narrative_text))

    if narrative_text:
        sop.heading1("Overview")
        set_outline_level(sop.doc.paragraphs[-1], 0)
        _narrative_body(sop, narrative_text)
    sop.heading1("Procedure")
    set_outline_level(sop.doc.paragraphs[-1], 0)
    for n, (step, result) in enumerate(zip(manifest.steps, step_results, strict=True), start=1):
        heading = step_heading(n, step)
        sop.heading2(heading)
        set_outline_level(sop.doc.paragraphs[-1], 1)
        _step_bullet(sop, step, result)
        if result.get("narration"):
            sop.bullet(f"Narration: {result['narration']}", sub=True)
        sop.image(step.screenshot, caption=heading)
    sop.heading1("Revision History")
    set_outline_level(sop.doc.paragraphs[-1], 0)
    sop.revision_history([(date, revision, "Initial generation", author)])
    enable_update_fields_on_open(sop.doc)
    out = sop.save()
    return out, sop.warnings
