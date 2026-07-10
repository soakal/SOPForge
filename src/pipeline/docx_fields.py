"""Real Word TOC field support, layered on top of the external SOP Factory 2
engine (`sop_lib.SOPBuilder`, CLAUDE.md: "extend it, do not rewrite it").

The engine's own `SOPBuilder.toc()` writes the table of contents as plain
static paragraphs — no `TOC` field code, so Word can never build page
numbers or navigation for it, and right-click "Update Field" has nothing to
act on. `SOPBuilder.heading1()`/`heading2()` likewise never set an outline
level, so even a real `TOC` field would have nothing to collect.

Rather than touch the external, unversioned engine file, this module drives
`SOPBuilder`'s own public methods for the visible formatting (so the VRSI
look is untouched) and layers real OOXML field/outline-level plumbing on
top via `sop.doc`, the `python-docx` Document the engine already exposes.
The static text `toc_lines()` produces becomes the field's *cached result*
(the runs Word shows before the first "Update Field") so every non-Word
consumer of the docx, and the pre-update view in Word itself, still show
exactly the same text as before this module existed.
"""

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_outline_level(paragraph, level):
    """Marks `paragraph` as outline level `level` (0-based, matching Word's
    "Level 1" = 0) by adding `<w:outlineLvl>` to its `pPr`. This is the
    signal a `TOC \\o \\u` field switch collects entries by — it works
    whether or not the paragraph carries a named Word heading style, so it
    doesn't require changing SOPBuilder's own heading1/heading2 formatting
    (which uses manually bolded/italicized runs, not built-in styles)."""
    pPr = paragraph._p.get_or_add_pPr()
    outline_lvl = OxmlElement("w:outlineLvl")
    outline_lvl.set(qn("w:val"), str(level))
    pPr.append(outline_lvl)


def add_toc_field(sop, cached_lines):
    """Replaces SOPBuilder.toc()'s plain-text table of contents with a real
    Word `TOC` field, using `cached_lines` (assembler.py's `toc_lines()`
    output — the same strings the old plain-text TOC showed) as the field's
    cached result. `\\o "1-2"` collects outline levels 0-1 (set_outline_level
    marks section headings 0, step headings 1); `\\h` makes entries
    hyperlinks; `\\z` hides tab leaders in web view; `\\u` collects by
    paragraph outline level in addition to style — the switch this whole
    mechanism relies on, since SOPBuilder's headings carry no named style."""
    sop.heading1("TABLE OF CONTENTS")

    paragraph = sop.doc.add_paragraph()
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    for line in cached_lines:
        sop.doc.add_paragraph(line)

    end_paragraph = sop.doc.add_paragraph()
    end_run = end_paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)

    sop.page_break()


def enable_update_fields_on_open(doc):
    """Sets `<w:updateFields w:val="true"/>` in word/settings.xml so Word
    refreshes the TOC field's page numbers/links automatically when the
    document is opened, instead of requiring a manual right-click ->
    "Update Field" (or F9) before it shows anything real."""
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
