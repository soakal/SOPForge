"""
sop_lib.py — VRSI SOP formatting helpers
========================================
Reusable building blocks for any SOP. The SOP Builder skill imports these so
every document gets the same VRSI look (header, footer, bullets, captions,
revision history table) regardless of subject matter.

Usage from a builder script:

    from sop_lib import SOPBuilder

    sop = SOPBuilder(
        template_docx="C:/.../template/SOP_TEMPLATE_WITH_PHOTOS.docx",
        output_docx="C:/.../output/My_SOP_v1.0.docx",
        active_dir="C:/.../active",
        revision="1.0",
        date="05/03/2026",
    )
    sop.title_page("BACKUP", "AND", "RESTORE",
                   subtitle="Trendsetter station — VRSI procedure",
                   doc_no="SOP-TS-001", author="BK (Brian Kalsic)")
    sop.toc([
        "1.  Overview",
        "2.  Backup", "      A)  ...", "      B)  ...",
        "3.  Restore", "      A)  ...",
        "4.  Revision History",
    ])
    sop.heading1("1.   Overview")
    sop.paragraph("This document covers ...")
    sop.heading1("2.   Backup")
    sop.heading2("A)   Reorient the Hidden Drive")
    sop.bullet("Power down the cabinet")
    sop.bullet_rich([("Click ", False), ("Backup", True), (" → ", False), ("System Backup", True)])
    sop.image("01_drive_layout.png", caption="Drive layout in Disk Management")
    sop.heading1("Revision History")
    sop.revision_history([("05/03/2026", "1.0", "First Draft", "BK")])
    sop.save()

The skill prompt should give Claude the freedom to call any of these helpers
in whatever order best matches the SOP's content.
"""

import io
import os
import shutil
from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple, Union

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageOps


def load_secret(key: str, base_dir: Union[str, Path]) -> str:
    """Load a secret by key: environment variable first, then KEY=VALUE lines
    in <base_dir>/secrets.txt (gitignored). Exits with a clear message if the
    secret cannot be found — builds must not silently embed a wrong value."""
    val = os.environ.get(key)
    if val:
        return val.strip()
    secrets_file = Path(base_dir) / "secrets.txt"
    if secrets_file.exists():
        for line in secrets_file.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                if k.strip() == key and v.strip():
                    return v.strip()
    raise SystemExit(
        f"Missing secret '{key}'. Set the {key} environment variable or add a "
        f"line '{key}=<value>' to {secrets_file} (kept out of version control)."
    )


class SOPBuilder:
    """Wraps a python-docx Document with VRSI-formatted helpers."""

    # ---- public API ----------------------------------------------------------

    def __init__(
        self,
        template_docx: Union[str, Path],
        output_docx: Union[str, Path],
        active_dir: Union[str, Path],
        revision: str = "1.0",
        date: str = "01/01/2026",
    ):
        self.template_docx = Path(template_docx)
        self.output_docx   = Path(output_docx)
        self.active_dir    = Path(active_dir)
        self.revision      = revision
        self.date          = date
        self.warnings: List[str] = []

        # Copy template (preserves header/footer/logo) then clear the body.
        self.output_docx.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(str(self.template_docx), str(self.output_docx))
        self.doc = Document(str(self.output_docx))
        self._clear_body()

        # Update header placeholders [X.X] / [MM/DD/YYYY] in-memory; we rewrite
        # the zip on save() to bake them in.
        self._header_revision = revision
        self._header_date     = date

    # ---- title / TOC ---------------------------------------------------------

    def title_page(
        self,
        line1: str,
        connector: str = "AND",
        line2: Optional[str] = None,
        *,
        subtitle: Optional[str] = None,
        doc_no: Optional[str] = None,
        author: Optional[str] = None,
    ):
        """Three-line title page with the VRSI 'AND' connector style."""
        for _ in range(2):
            self.doc.add_paragraph("")
        self._centered_run(line1, size=Pt(26), bold=True, underline=True)
        if line2:
            self._centered_run(connector, size=Pt(20), italic=True)
            self._centered_run(line2, size=Pt(26), bold=True, underline=True)
        if subtitle:
            self.doc.add_paragraph("")
            self._centered_run(subtitle, size=Pt(13), italic=True)
        for _ in range(3):
            self.doc.add_paragraph("")
        if doc_no:
            self._centered_text(f"Document No: {doc_no}")
        self._centered_text(f"Revision {self.revision}  —  Released {self.date}")
        if author:
            self._centered_text(f"Author: {author}")
        self.page_break()

    def toc(self, lines: Iterable[str]):
        """Plain Table of Contents (text — no auto-numbered TOC field)."""
        self.heading1("TABLE OF CONTENTS")
        for line in lines:
            self.doc.add_paragraph(line)
        self.page_break()

    # ---- structure -----------------------------------------------------------

    def heading1(self, text: str):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)

    def heading2(self, text: str):
        """Bold italic — VRSI subsection style."""
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.italic = True
        r.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(3)

    def paragraph(self, text: str):
        self.doc.add_paragraph(text)

    def page_break(self):
        self.doc.add_page_break()

    # ---- bullets -------------------------------------------------------------

    def bullet(self, text: str, *, sub: bool = False, system_only: Optional[str] = None):
        """Plain bullet. system_only adds a red bold-italic '— Ford Only' style flag."""
        p = self._bullet_paragraph(sub)
        p.add_run(text)
        if system_only:
            f = p.add_run(f"  — {system_only}")
            f.bold = True
            f.italic = True
            f.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        return p

    def bullet_rich(
        self,
        parts: Sequence[Tuple[str, Union[bool, str]]],
        *,
        sub: bool = False,
        system_only: Optional[str] = None,
    ):
        """Bullet built from (text, style) tuples.

        style: False = plain, True or 'bold' = bold (UI element names / typed
        values), 'code' = monospace (file paths, commands, registry keys)."""
        p = self._bullet_paragraph(sub)
        for text, style in parts:
            r = p.add_run(text)
            if style == "code":
                r.font.name = "Consolas"
                r.font.size = Pt(10)
                rPr = r._element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.append(rFonts)
                for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                    rFonts.set(qn(attr), "Consolas")
            else:
                r.bold = bool(style)
        if system_only:
            f = p.add_run(f"  — {system_only}")
            f.bold = True
            f.italic = True
            f.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        return p

    # ---- images --------------------------------------------------------------

    def image(self, filename: str, *, caption: Optional[str] = None,
              width_in: float = 5.0, max_px: int = 1600, jpeg_quality: int = 88):
        """Insert numbered image from active/ at the current cursor position."""
        path = self.active_dir / filename
        if not path.exists():
            self.doc.add_paragraph(f"[IMAGE NOT FOUND: {filename}]")
            self.warnings.append(f"image not found: {filename} (looked in {self.active_dir})")
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            data = self._prep_image(path, max_px=max_px, jpeg_quality=jpeg_quality)
            run.add_picture(io.BytesIO(data), width=Inches(width_in))
        except Exception as e:
            p.clear()
            p.add_run(f"[Could not insert {filename}: {e}]")
        if caption:
            cap = self.doc.add_paragraph(caption)
            cap.runs[0].italic = True
            cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph("")

    def placeholder(self, description: str):
        """[IMAGE PLACEHOLDER: ...] line — for steps where no image exists yet."""
        p = self.doc.add_paragraph(f"[IMAGE PLACEHOLDER: {description}]")
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---- tables --------------------------------------------------------------

    def revision_history(self, rows: Sequence[Tuple[str, str, str, str]]):
        """4-column Date | Version | Comments | Authors table with cell borders."""
        table = self.doc.add_table(rows=1 + len(rows), cols=4)
        self._add_cell_borders(table)
        hdr = table.rows[0].cells
        for i, t in enumerate(("Date", "Version", "Comments", "Authors")):
            hdr[i].text = t
            hdr[i].paragraphs[0].runs[0].bold = True
        for r_idx, row in enumerate(rows, start=1):
            cells = table.rows[r_idx].cells
            for i, val in enumerate(row):
                cells[i].text = val

    def scenario_table(self, rows: Sequence[Tuple[str, str]],
                       headers: Tuple[str, str] = ("Scenario", "Example Value")):
        """2-column scenario table (e.g., task naming examples)."""
        table = self.doc.add_table(rows=1 + len(rows), cols=2)
        self._add_cell_borders(table)
        hdr = table.rows[0].cells
        for i, t in enumerate(headers):
            hdr[i].text = t
            hdr[i].paragraphs[0].runs[0].bold = True
        for r_idx, row in enumerate(rows, start=1):
            cells = table.rows[r_idx].cells
            for i, val in enumerate(row):
                cells[i].text = val

    # ---- save ----------------------------------------------------------------

    def save(self):
        """Save and patch [X.X]/[MM/DD/YYYY] header placeholders in the zip.
        Prints any accumulated warnings (missing images, unpatched headers)."""
        self.doc.save(str(self.output_docx))
        self._patch_header_placeholders()
        self._check_headers_patched()
        for w in self.warnings:
            print(f"WARN: {w}")
        return self.output_docx

    def _check_headers_patched(self):
        """Warn if [X.X] / [MM/DD/YYYY] survived the header patch (e.g. the
        placeholder was split across XML runs in the template)."""
        import zipfile
        with zipfile.ZipFile(self.output_docx, "r") as z:
            for item in z.namelist():
                if item.startswith("word/header") and item.endswith(".xml"):
                    txt = z.read(item).decode("utf-8", errors="ignore")
                    for ph in ("[X.X]", "[MM/DD/YYYY]"):
                        if ph in txt:
                            self.warnings.append(
                                f"header placeholder {ph} not patched in {item}")

    # ---- private -------------------------------------------------------------

    def _clear_body(self):
        body = self.doc._element.body
        for child in list(body):
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)

    def _centered_run(self, text, *, size, bold=False, italic=False, underline=False):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.size = size
        r.bold = bold
        r.italic = italic
        r.underline = underline

    def _centered_text(self, text):
        p = self.doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _bullet_paragraph(self, sub: bool):
        p = self.doc.add_paragraph()
        if sub:
            p.paragraph_format.left_indent = Inches(0.6)
            p.add_run("○  ")
        else:
            p.paragraph_format.left_indent = Inches(0.25)
            p.add_run("•  ")
        return p

    def _prep_image(self, path: Path, *, max_px: int, jpeg_quality: int) -> bytes:
        img = Image.open(str(path))
        try:
            img = ImageOps.exif_transpose(img)
        except Exception:
            pass
        w, h = img.size
        if max(w, h) > max_px:
            s = max_px / max(w, h)
            img = img.resize((int(w * s), int(h * s)), Image.LANCZOS)
        buf = io.BytesIO()
        # Keep PNG screenshots lossless (JPEG smears UI text); photos stay JPEG.
        if path.suffix.lower() == ".png":
            if img.mode not in ("RGB", "RGBA", "P", "L"):
                img = img.convert("RGBA")
            img.save(buf, format="PNG", optimize=True)
        else:
            img.convert("RGB").save(buf, format="JPEG", quality=jpeg_quality, optimize=True)
        return buf.getvalue()

    @staticmethod
    def _add_cell_borders(table):
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top", "left", "bottom", "right"):
                    b = OxmlElement(f"w:{side}")
                    b.set(qn("w:val"), "single")
                    b.set(qn("w:sz"), "4")
                    b.set(qn("w:color"), "000000")
                    tcBorders.append(b)
                tcPr.append(tcBorders)

    def _patch_header_placeholders(self):
        """Rewrite the docx zip, replacing header placeholder strings."""
        import zipfile, tempfile, os
        src = self.output_docx
        tmp_fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=str(src.parent))
        os.close(tmp_fd)
        with zipfile.ZipFile(src, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.namelist():
                data = zin.read(item)
                if item.startswith("word/header") and item.endswith(".xml"):
                    txt = data.decode("utf-8")
                    txt = (txt
                           .replace("[X.X]", self._header_revision)
                           .replace("[MM/DD/YYYY]", self._header_date))
                    data = txt.encode("utf-8")
                zout.writestr(item, data)
        shutil.move(tmp_path, str(src))
